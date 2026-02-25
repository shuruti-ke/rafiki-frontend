"""
Parse payroll files (CSV / Excel / PDF / Word) and match employees to users_legacy.

Supported:
- .csv
- .xlsx (openpyxl)
- .xls  (xlrd)
- .docx (python-docx)
- .pdf  (pdfplumber) best-effort table extraction
"""
import csv
import io
import logging
import re
from typing import Optional, Tuple

from sqlalchemy.orm import Session
from app.models.user import User

logger = logging.getLogger(__name__)

# ---- Optional imports (keep server from crashing if missing) ----
try:
    import openpyxl  # .xlsx
except ImportError:
    openpyxl = None

try:
    import xlrd  # .xls (legacy Excel)
except ImportError:
    xlrd = None

try:
    import docx  # python-docx for .docx
except ImportError:
    docx = None

try:
    import pdfplumber  # for .pdf table extraction
except ImportError:
    pdfplumber = None


def _normalize_name(name: str) -> str:
    return " ".join(name.strip().lower().split())


def _parse_number(val) -> float:
    """Parse numbers like '1,234.56' or '1 234.56' safely."""
    if isinstance(val, (int, float)):
        return float(val)
    try:
        cleaned = str(val).replace(",", "").replace(" ", "").strip()
        # strip currency symbols
        cleaned = re.sub(r"[^\d.\-]", "", cleaned)
        return float(cleaned) if cleaned else 0.0
    except (ValueError, TypeError):
        return 0.0


def _match_employees(entries: list[dict], db: Session, org_id) -> list[dict]:
    """
    For each parsed entry, try to match employee_name to users_legacy.name
    within the same org. Adds matched_user_id (str | None) to each entry.
    """
    users = db.query(User).filter(User.org_id == org_id, User.is_active == True).all()
    name_map = {}
    for u in users:
        if u.name:
            name_map[_normalize_name(u.name)] = str(u.user_id)

    for entry in entries:
        norm = _normalize_name(entry["employee_name"])
        entry["matched_user_id"] = name_map.get(norm)

    return entries


def _build_summary(entries: list[dict]) -> dict:
    total_gross = sum(e["gross_salary"] for e in entries)
    total_deductions = sum(e["deductions"] for e in entries)
    total_net = sum(e["net_salary"] for e in entries)
    unmatched = [e["employee_name"] for e in entries if not e.get("matched_user_id")]

    expected_net = total_gross - total_deductions
    # Tolerance: 1% of total gross or 1.0, whichever is larger — handles rounding & other additions
    tolerance = max(1.0, total_gross * 0.01)
    reconciled = abs(expected_net - total_net) < tolerance

    return {
        "entries": entries,
        "total_gross": round(total_gross, 2),
        "total_deductions": round(total_deductions, 2),
        "total_net": round(total_net, 2),
        "employee_count": len(entries),
        "matched_count": sum(1 for e in entries if e.get("matched_user_id")),
        "unmatched_names": unmatched,
        "reconciled": reconciled,
    }


# ----------------------------
# Column mapping helpers
# ----------------------------

def _normalize_header(h) -> str:
    if h is None:
        return ""
    return str(h).strip().lower().replace(" ", "_")


def _row_to_entry(normed: dict) -> Optional[dict]:
    """
    Convert a normalized row dict to a payroll entry.
    Expected columns (case-insensitive):
      employee_name (or name), gross_salary (or gross), deductions, net_salary (or net)
    """
    name = normed.get("employee_name") or normed.get("name") or ""
    if isinstance(name, (int, float)):
        name = str(name)
    name = str(name).strip()
    if not name:
        return None

    gross = _parse_number(
        normed.get("gross_salary") or normed.get("gross") or normed.get("gross_pay") or 0
    )
    # Sum individual deduction columns if a single "deductions" column isn't present
    if normed.get("deductions"):
        deductions = _parse_number(normed["deductions"])
    else:
        deduction_keys = {"paye", "nssf", "shif", "nhdf", "nhif", "levy", "loan", "deduction"}
        deductions = sum(
            abs(_parse_number(v))
            for k, v in normed.items()
            if any(dk in k for dk in deduction_keys) and v not in (None, "")
        )
    net = _parse_number(
        normed.get("net_salary") or normed.get("net") or normed.get("net_pay") or 0
    )

    known_keys = {"employee_name", "name", "gross_salary", "gross", "deductions", "net_salary", "net"}
    details = {k: (str(v).strip() if v is not None else "") for k, v in normed.items() if k not in known_keys}
    details = {k: v for k, v in details.items() if v}

    return {
        "employee_name": name,
        "gross_salary": gross,
        "deductions": deductions,
        "net_salary": net,
        "details": details if details else None,
    }


# ----------------------------
# CSV
# ----------------------------

def parse_csv_payroll(content: bytes, db: Session, org_id) -> dict:
    text = content.decode("utf-8-sig", errors="ignore")
    reader = csv.DictReader(io.StringIO(text))

    if not reader.fieldnames:
        return {"entries": [], "total_gross": 0, "total_deductions": 0, "total_net": 0, "unmatched_names": []}

    col_map = {col: _normalize_header(col) for col in reader.fieldnames}

    entries = []
    for row in reader:
        normed = {col_map.get(k, k): (v.strip() if isinstance(v, str) else v) for k, v in row.items()}
        entry = _row_to_entry(normed)
        if entry:
            entries.append(entry)

    entries = _match_employees(entries, db, org_id)
    return _build_summary(entries)


# ----------------------------
# Excel .xlsx
# ----------------------------

def _find_header_row(rows: list) -> int:
    """
    Find the best header row index in the first 15 rows.
    Strategy: pick the row with the most non-empty string cells.
    This skips leading empty rows, title rows, and merged group headers.
    Also look for a row that contains keywords like 'name', 'gross', 'net'.
    """
    PAYROLL_KEYWORDS = {"name", "gross", "net", "salary", "pay", "employee"}
    best_idx = None
    best_score = -1

    for i, row in enumerate(rows[:15]):
        str_cells = [c for c in row if c is not None and not isinstance(c, (int, float)) and str(c).strip()]
        keyword_hits = sum(
            1 for c in str_cells if any(kw in str(c).lower() for kw in PAYROLL_KEYWORDS)
        )
        # Score: keyword matches are worth 3x more than plain string cells
        score = keyword_hits * 3 + len(str_cells)
        if score > best_score:
            best_score = score
            best_idx = i

    return best_idx if (best_idx is not None and best_score > 0) else 0


def parse_xlsx_payroll(content: bytes, db: Session, org_id) -> dict:
    if openpyxl is None:
        logger.error("openpyxl not installed — cannot parse .xlsx files")
        return {"entries": [], "total_gross": 0, "total_deductions": 0, "total_net": 0, "unmatched_names": [], "error": "openpyxl not installed"}

    wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))
    wb.close()

    if not rows:
        return {"entries": [], "total_gross": 0, "total_deductions": 0, "total_net": 0, "unmatched_names": []}

    header_idx = _find_header_row(rows)
    raw_headers = rows[header_idx]

    # Build column name list: blank headers get a positional fallback name
    headers = []
    for i, h in enumerate(raw_headers):
        normalized = _normalize_header(h)
        headers.append(normalized if normalized else f"col_{i}")

    entries = []
    for row in rows[header_idx + 1:]:
        # Skip entirely empty rows (totals/footer rows are often not fully empty, but skip if all None)
        if all(c is None for c in row):
            continue
        normed = {headers[i]: row[i] for i in range(len(headers)) if i < len(row)}
        entry = _row_to_entry(normed)
        if entry:
            entries.append(entry)

    entries = _match_employees(entries, db, org_id)
    return _build_summary(entries)


# ----------------------------
# Excel .xls (legacy)
# ----------------------------

def parse_xls_payroll(content: bytes, db: Session, org_id) -> dict:
    if xlrd is None:
        logger.error("xlrd not installed — cannot parse .xls files")
        return {"entries": [], "total_gross": 0, "total_deductions": 0, "total_net": 0, "unmatched_names": [], "error": "xlrd not installed"}

    book = xlrd.open_workbook(file_contents=content)
    sheet = book.sheet_by_index(0)

    if sheet.nrows < 1:
        return {"entries": [], "total_gross": 0, "total_deductions": 0, "total_net": 0, "unmatched_names": []}

    headers = [_normalize_header(sheet.cell_value(0, c)) for c in range(sheet.ncols)]
    entries = []
    for r in range(1, sheet.nrows):
        normed = {headers[c]: sheet.cell_value(r, c) for c in range(sheet.ncols)}
        entry = _row_to_entry(normed)
        if entry:
            entries.append(entry)

    entries = _match_employees(entries, db, org_id)
    return _build_summary(entries)


# ----------------------------
# Word .docx
# ----------------------------

def parse_docx_payroll(content: bytes, db: Session, org_id) -> dict:
    if docx is None:
        logger.error("python-docx not installed — cannot parse .docx files")
        return {"entries": [], "total_gross": 0, "total_deductions": 0, "total_net": 0, "unmatched_names": [], "error": "python-docx not installed"}

    d = docx.Document(io.BytesIO(content))
    entries = []

    # Prefer tables (most payroll Word docs use tables)
    for table in d.tables:
        rows = table.rows
        if not rows:
            continue
        headers = [_normalize_header(cell.text) for cell in rows[0].cells]
        if not any(headers):
            continue

        for r in rows[1:]:
            values = [cell.text.strip() for cell in r.cells]
            normed = {headers[i]: values[i] for i in range(min(len(headers), len(values)))}
            entry = _row_to_entry(normed)
            if entry:
                entries.append(entry)

    # If no tables found, try line-based parsing (best effort)
    if not entries:
        full_text = "\n".join(p.text for p in d.paragraphs if p.text.strip())
        entries = _best_effort_parse_lines(full_text)

    entries = _match_employees(entries, db, org_id)
    return _build_summary(entries)


# ----------------------------
# PDF
# ----------------------------

def parse_pdf_payroll(content: bytes, db: Session, org_id) -> dict:
    if pdfplumber is None:
        logger.error("pdfplumber not installed — cannot parse .pdf files")
        return {"entries": [], "total_gross": 0, "total_deductions": 0, "total_net": 0, "unmatched_names": [], "error": "pdfplumber not installed"}

    entries = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            # Try tables first
            tables = page.extract_tables() or []
            for tbl in tables:
                if not tbl or len(tbl) < 2:
                    continue
                headers = [_normalize_header(h) for h in tbl[0]]
                if not any(headers):
                    continue
                for row in tbl[1:]:
                    if not row:
                        continue
                    normed = {headers[i]: (row[i] if i < len(row) else None) for i in range(len(headers))}
                    entry = _row_to_entry(normed)
                    if entry:
                        entries.append(entry)

            # If no tables, fall back to text lines (best effort)
            if not tables:
                text = page.extract_text() or ""
                if text.strip():
                    entries.extend(_best_effort_parse_lines(text))

    entries = _match_employees(entries, db, org_id)
    return _build_summary(entries)


# ----------------------------
# Best-effort fallback parsing
# ----------------------------

def _best_effort_parse_lines(text: str) -> list[dict]:
    """
    Tries to parse lines that look like:
      "Jane Doe, 5000, 500, 4500"
      "Jane Doe 5000 500 4500"
    This is fallback for PDF/Word text when tables are not extractable.
    """
    entries = []
    for line in text.splitlines():
        line = line.strip()
        if not line or len(line) < 8:
            continue

        # Try comma-separated
        if line.count(",") >= 3:
            parts = [p.strip() for p in line.split(",")]
        else:
            # Try whitespace split with at least 4 components
            parts = re.split(r"\s{2,}|\t+", line)
            parts = [p.strip() for p in parts if p.strip()]

        if len(parts) < 4:
            continue

        # Assume: name, gross, deductions, net
        name = parts[0]
        gross = _parse_number(parts[1])
        deductions = _parse_number(parts[2])
        net = _parse_number(parts[3])

        if not name or (gross == 0 and deductions == 0 and net == 0):
            continue

        entries.append({
            "employee_name": name,
            "gross_salary": gross,
            "deductions": deductions,
            "net_salary": net,
            "details": {"source": "best_effort_line_parse", "raw": line},
        })

    return entries


# ----------------------------
# Unified entry point
# ----------------------------

def parse_payroll_file(filename: str, content: bytes, db: Session, org_id) -> dict:
    """
    Decide parser based on filename extension.
    """
    ext = (filename.rsplit(".", 1)[-1].lower() if "." in filename else "").strip()

    if ext == "csv":
        return parse_csv_payroll(content, db, org_id)
    if ext == "xlsx":
        return parse_xlsx_payroll(content, db, org_id)
    if ext == "xls":
        return parse_xls_payroll(content, db, org_id)
    if ext == "docx":
        return parse_docx_payroll(content, db, org_id)
    if ext == "pdf":
        return parse_pdf_payroll(content, db, org_id)

    return {
        "entries": [],
        "total_gross": 0,
        "total_deductions": 0,
        "total_net": 0,
        "unmatched_names": [],
        "error": f"Unsupported file type: .{ext or 'unknown'}"
    }
